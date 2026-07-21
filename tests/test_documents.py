"""Phase 10 — /v1/documents, as executable claims.

Headlines:
  - the replace test: a shorter revision must leave NO orphaned chunks of
    the old version (delete-then-upsert, not upsert-and-hope)
  - the tenant test: the form cannot say whose corpus changes — vikram's
    token uploads into vikram's corpus no matter what anyone types
  - the fusion regression: two documents may both have a chunk 0; fusion
    must never merge them (the fixture's global enumerate was load-bearing)
"""

import io                               # stdlib — in-memory upload bodies

import pytest                           # 3rd-party: pytest — fixtures

from fastapi.testclient import TestClient  # 3rd-party: fastapi (submodule) — drives the app

from app.auth import mint               # local — app/auth.py
from app.config import get_settings      # local — app/config.py
from app.ingest import Chunk            # local — app/ingest/
from app.ingest.loaders import _looks_like_heading  # local — the PDF heuristic
from app.main import app                # local — app/main.py
from app.retrieval.hybrid import Hit, fuse_rrf  # local — app/retrieval/hybrid.py
from app.store.base import Gate         # local — app/store/base.py


def auth(tenant: str = "asha", groups: tuple[str, ...] = ("agent",)) -> dict:
    token = mint(tenant, list(groups), secret=get_settings().auth_jwt_secret)
    return {"Authorization": f"Bearer {token}"}


TODAY = __import__("datetime").date(2026, 7, 21)


def _visible_titles(client, tenant: str) -> set[str]:
    gate = Gate(tenant, frozenset({"customer", "agent"}), TODAY)
    return {c.doc_title for c in client.app.state.store.visible_chunks(gate)}


# -----------------------------------------------------------------------------
# A minimal but VALID single-page PDF, built by hand so the test suite needs
# no PDF-writing dependency. One text line per input string; pypdf reads it.
# -----------------------------------------------------------------------------
def tiny_pdf(lines: list[str]) -> bytes:
    body = ["BT /F1 12 Tf 50 750 Td 16 TL"]
    for ln in lines:
        esc = ln.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        body.append(f"({esc}) Tj T*")
    body.append("ET")
    stream = "\n".join(body).encode("latin-1")

    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objs) + 1}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
    return bytes(out)


def _upload(client, *, filename: str, data: bytes, title: str, headers: dict, **form):
    return client.post(
        "/v1/documents",
        files={"file": (filename, io.BytesIO(data), "application/octet-stream")},
        data={"title": title, **form},
        headers=headers,
    )


# =============================================================================
# The boundary gates
# =============================================================================
def test_customers_cannot_write_the_corpus() -> None:
    with TestClient(app) as client:
        r = _upload(
            client, filename="a.md", data=b"# T\n\nbody", title="Customer Upload",
            headers=auth(groups=("customer",)),
        )
    assert r.status_code == 403
    assert r.json()["error"]["code"] == "forbidden"


def test_unsupported_types_are_415_not_500() -> None:
    with TestClient(app) as client:
        r = _upload(
            client, filename="malware.exe", data=b"MZ....", title="Nice Try",
            headers=auth(),
        )
    assert r.status_code == 415


def test_oversized_uploads_are_413_before_parsing() -> None:
    small_cap = get_settings().model_copy(update={"max_upload_bytes": 64})
    app.dependency_overrides[get_settings] = lambda: small_cap
    try:
        with TestClient(app) as client:
            r = _upload(
                client, filename="big.md", data=b"# T\n\n" + b"x" * 200,
                title="Too Big", headers=auth(),
            )
    finally:
        app.dependency_overrides.pop(get_settings, None)
    assert r.status_code == 413


def test_corrupt_pdf_is_the_clients_422() -> None:
    with TestClient(app) as client:
        r = _upload(
            client, filename="broken.pdf", data=b"%PDF-1.4 then garbage",
            title="Broken PDF", headers=auth(),
        )
    assert r.status_code == 422
    assert r.json()["error"]["code"] in ("unparseable", "empty_document")


# =============================================================================
# The upload -> answerable loop, and whose corpus changes
# =============================================================================
_ADDENDUM = """# Roadside Addendum

## 1. Towing

Towing to the nearest approved garage is arranged within ninety (90) minutes
of the assistance call being logged.
"""


def test_markdown_upload_is_immediately_answerable() -> None:
    with TestClient(app) as client:
        up = _upload(
            client, filename="addendum.md", data=_ADDENDUM.encode(),
            title="Asha Rao — Roadside Addendum", headers=auth(),
        )
        assert up.status_code == 201
        assert up.json()["chunks"] > 0
        assert up.json()["tenant_id"] == "asha"

        # Same store the retriever reads — no reload, no re-boot. Anchored
        # phrasing (the document's own words) so the reranker clears the gate.
        r = client.post(
            "/v1/ask",
            json={"question": "how quickly is towing to the nearest approved garage arranged?"},
            headers=auth(),
        )
        assert '"type": "sources"' in r.text or '"type":"sources"' in r.text
        assert "Roadside Addendum" in r.text


def test_tenant_comes_from_the_token_not_the_form() -> None:
    """Vikram's agent uploads. Nothing in the form said 'vikram' — and nothing
    in the form COULD say 'asha': the corpus that changes belongs to the
    verified token, and asha's view never contains the document."""
    with TestClient(app) as client:
        up = _upload(
            client, filename="v.md", data=b"# Garage List\n\nApproved garages for Vikram.",
            title="Vikram — Garage List", headers=auth("vikram"),
        )
        assert up.status_code == 201
        assert up.json()["tenant_id"] == "vikram"

        assert "Vikram — Garage List" in _visible_titles(client, "vikram")
        assert "Vikram — Garage List" not in _visible_titles(client, "asha")


def test_replacing_with_a_shorter_revision_leaves_no_orphans() -> None:
    long_doc = "# Kit\n\n" + "\n\n".join(
        f"## Part {i}\n\nBody text for part {i}, marker LONGVER-{i}." for i in range(1, 9)
    )
    short_doc = "# Kit\n\n## Part 1\n\nBody text, marker SHORTVER-1."

    with TestClient(app) as client:
        first = _upload(client, filename="kit.md", data=long_doc.encode(),
                        title="Asha Rao — Kit", headers=auth())
        n_long = first.json()["chunks"]

        second = _upload(client, filename="kit.md", data=short_doc.encode(),
                         title="Asha Rao — Kit", headers=auth())
        assert second.status_code == 201
        assert second.json()["replaced_chunks"] == n_long, "the whole old version went first"

        gate = Gate("asha", frozenset({"agent"}), TODAY)
        kit_chunks = [
            c for c in client.app.state.store.visible_chunks(gate)
            if c.doc_title == "Asha Rao — Kit"
        ]
        assert len(kit_chunks) == second.json()["chunks"]
        # The orphan test proper: no trace of the long version's tail.
        assert not any("LONGVER" in c.text for c in kit_chunks)


# =============================================================================
# Formats
# =============================================================================
def test_pdf_upload_infers_headings() -> None:
    pdf = tiny_pdf([
        "1. Cover",
        "Own damage cover applies to the insured vehicle.",
        "GENERAL EXCLUSIONS",
        "Wear and tear is not covered by this addendum.",
    ])
    with TestClient(app) as client:
        up = _upload(client, filename="cover.pdf", data=pdf,
                     title="Asha Rao — PDF Addendum", headers=auth())
        assert up.status_code == 201

        gate = Gate("asha", frozenset({"agent"}), TODAY)
        headings = {
            c.heading for c in client.app.state.store.visible_chunks(gate)
            if c.doc_title == "Asha Rao — PDF Addendum"
        }
    assert "1. Cover" in headings
    assert "GENERAL EXCLUSIONS" in headings


def test_docx_upload_extracts_styles_and_tables() -> None:
    from docx import Document           # 3rd-party: python-docx — writer, in tests only

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Approved Repairers", level=1)
    doc.add_paragraph("Repairs must be carried out by an approved repairer.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text = "City", "Repairer"
    t.cell(1, 0).text, t.cell(1, 1).text = "Pune", "Deccan Motors"
    doc.save(buf)

    with TestClient(app) as client:
        up = _upload(client, filename="repairers.docx", data=buf.getvalue(),
                     title="Asha Rao — Repairers", headers=auth())
        assert up.status_code == 201

        gate = Gate("asha", frozenset({"agent"}), TODAY)
        chunks = [
            c for c in client.app.state.store.visible_chunks(gate)
            if c.doc_title == "Asha Rao — Repairers"
        ]
    assert any(c.heading == "Approved Repairers" for c in chunks)
    assert any(c.is_table and "Deccan Motors" in c.text for c in chunks)


def test_heading_heuristic_edges() -> None:
    assert _looks_like_heading("7. Limit of Liability")
    assert _looks_like_heading("2.1 Instalments")
    assert _looks_like_heading("GENERAL EXCLUSIONS")
    assert not _looks_like_heading("This is an ordinary sentence about cover.")
    assert not _looks_like_heading("WEAR AND TEAR IS NOT COVERED, SEE SECTION 4.")


# =============================================================================
# The fusion regression the fixture was hiding
# =============================================================================
def test_fusion_never_merges_same_index_across_documents() -> None:
    a = Chunk(doc_title="Doc A", heading="h", text="alpha", parent_text="alpha", chunk_index=0)
    b = Chunk(doc_title="Doc B", heading="h", text="beta", parent_text="beta", chunk_index=0)

    fused = fuse_rrf([[Hit(a, 0, 1.0)], [Hit(b, 0, 9.9)]], k=10)

    texts = {h.chunk.text for h in fused}
    assert texts == {"alpha", "beta"}, (
        "two documents' chunk 0 must fuse as TWO results — bare chunk_index "
        "keying silently merged them (the fixture's global enumerate hid this)"
    )
