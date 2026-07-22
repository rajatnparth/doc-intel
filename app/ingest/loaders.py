"""Step 1: read a document into TEXT + STRUCTURE.

The most important idea in this file: a loader's job is not "get the text out".
It is "get the text out WITHOUT throwing away the structure the author put in".

Naive loaders flatten a document to a wall of characters. Then the chunker has
nothing to cut on except character count — and character-count cutting is the
failure the whole section is about. Structure preserved here is structure the
chunker can use later. Structure lost here is lost forever.
"""

from __future__ import annotations      # stdlib (special) — lets annotations be lazy strings,
                                        #   so `list[Section]` works and forward refs are free.
                                        #   Must be the first statement in the file.

import re                               # stdlib — regex: split on headings, find tables
from dataclasses import dataclass, field  # stdlib — @dataclass for Section; field() for the
                                        #   mutable default (default_factory=list)


@dataclass
class Section:
    """One coherent unit of a document — a heading and the body beneath it.

    This is the SEAM between loading and chunking. The loader produces Sections;
    the chunker consumes them. Notice it carries `heading` separately from
    `body`: that heading is the context a fragment needs to stay meaningful
    (section 3.1, 'contextual retrieval').
    """

    doc_title: str
    heading: str            # e.g. "2. Payment Terms"
    level: int              # 1 for #, 2 for ##, ...
    body: str
    # Blocks we must NOT let the chunker split by character count — tables, code.
    # Extracted whole, here, before any splitting can touch them.
    atomic_blocks: list[str] = field(default_factory=list)


# A markdown table is a run of consecutive lines that all contain a pipe.
# This regex finds such runs so we can pull them out as atomic units.
_TABLE_RE = re.compile(r"(?:^\|.*\|\s*$\n?)+", re.MULTILINE)


def load_markdown(text: str, *, doc_title: str) -> list[Section]:
    """Split markdown into Sections on headings, keeping tables intact.

    Why markdown first: contracts, invoices, wikis and exported PDFs all reduce
    cleanly to markdown, and markdown makes structure explicit (`##`, `|`). Real
    PDFs are messier — see load_pdf below — but the SHAPE of the answer is the
    same, so we learn it on the clean case.
    """
    sections: list[Section] = []

    # Split on headings but KEEP them (the capturing group in re.split keeps the
    # delimiter). Every odd element is a heading line, every even one its body.
    parts = re.split(r"^(#{1,6}\s+.*)$", text, flags=re.MULTILINE)

    # parts[0] is any preamble before the first heading; usually empty here.
    i = 1
    while i < len(parts):
        heading_line = parts[i].strip()
        body = parts[i + 1] if i + 1 < len(parts) else ""
        i += 2

        level = len(heading_line) - len(heading_line.lstrip("#"))
        heading_text = heading_line.lstrip("#").strip()

        # Pull tables OUT of the body before anyone can character-chunk them.
        # We replace each table with a placeholder so the prose flows, and store
        # the raw table in atomic_blocks to be emitted as its own chunk.
        atomic: list[str] = []

        def _extract(match: re.Match[str]) -> str:
            atomic.append(match.group(0).strip())
            return f"\n[TABLE {len(atomic)} EXTRACTED]\n"

        prose_body = _TABLE_RE.sub(_extract, body).strip()

        sections.append(
            Section(
                doc_title=doc_title,
                heading=heading_text,
                level=level,
                body=prose_body,
                atomic_blocks=atomic,
            )
        )

    return sections


# -----------------------------------------------------------------------------
# PDF — text layer + heading INFERENCE (phase 10).
#
# Real PDFs do not hand you headings; they hand you positioned text runs. A
# text-layer loader can still infer structure from two honest signals:
#   - numbered headings ("7. Limit of Liability", "2.1 Instalments")
#   - short shouting lines (ALL-CAPS, no terminal punctuation)
# Both are HEURISTICS and say so. Font-size/layout analysis, multi-column
# reading order, scanned pages — that is Azure Document Intelligence behind
# the same parse seam, on the private side of the split.
# -----------------------------------------------------------------------------
_NUMBERED_HEADING_RE = re.compile(r"^\d+(\.\d+)*[.)]?\s+\S")


def _looks_like_heading(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 80:
        return False
    if _NUMBERED_HEADING_RE.match(line):
        return True
    # Short, shouty, and not a sentence: "GENERAL EXCLUSIONS"
    return line.isupper() and not line.endswith((".", ":", ";", ","))


def load_pdf(data: bytes, *, doc_title: str) -> list[Section]:
    """Read a PDF (as bytes — uploads never touch the filesystem) into
    Sections: extract the text layer, strip page furniture, infer headings."""
    import io                            # stdlib — wrap the upload bytes for pypdf

    try:
        from pypdf import PdfReader     # 3rd-party: pypdf — pure-Python PDF reader.
                                        #   Imported LAZILY (inside the function, not at top)
                                        #   so markdown-only users never need it installed.
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("pip install pypdf to read PDFs") from e

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]

    # Strip lines that appear on nearly every page — they are headers/footers.
    # Left in, every chunk contains "CONFIDENTIAL - Page 12 of 40", and every
    # vector drifts slightly toward every other. This is not optional cleanup.
    line_counts: dict[str, int] = {}
    for pg in pages:
        for line in {ln.strip() for ln in pg.splitlines() if ln.strip()}:
            line_counts[line] = line_counts.get(line, 0) + 1
    threshold = max(2, int(len(pages) * 0.6))
    furniture = {ln for ln, c in line_counts.items() if c >= threshold}

    lines = [
        ln for pg in pages for ln in pg.splitlines() if ln.strip() not in furniture
    ]

    # Walk the lines, opening a new Section at every inferred heading. Text
    # before the first heading gets the document title as its heading — same
    # convention as load_markdown's preamble handling.
    sections: list[Section] = []
    heading, body = doc_title, []
    for ln in lines:
        if _looks_like_heading(ln):
            if body and any(b.strip() for b in body):
                sections.append(
                    Section(doc_title=doc_title, heading=heading, level=1, body="\n".join(body).strip())
                )
            heading, body = ln.strip(), []
        else:
            body.append(ln)
    if body and any(b.strip() for b in body):
        sections.append(
            Section(doc_title=doc_title, heading=heading, level=1, body="\n".join(body).strip())
        )
    return sections


def load_docx(data: bytes, *, doc_title: str) -> list[Section]:
    """Read a .docx into Sections. The EASY format: Word styles carry real
    heading levels, so this is extraction, not inference. Tables come out as
    markdown pipe rows into atomic_blocks — same non-splittable treatment as
    markdown tables."""
    import io                            # stdlib — wrap the upload bytes

    try:
        from docx import Document        # 3rd-party: python-docx — .docx reader
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("pip install python-docx to read .docx files") from e

    doc = Document(io.BytesIO(data))
    sections: list[Section] = []
    heading, level, body = doc_title, 1, []

    def flush() -> None:
        if body and any(b.strip() for b in body):
            sections.append(
                Section(doc_title=doc_title, heading=heading, level=level, body="\n".join(body).strip())
            )

    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") and para.text.strip():
            flush()
            heading, body = para.text.strip(), []
            level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
        elif para.text.strip():
            body.append(para.text)
    flush()

    # Tables: python-docx exposes them separately from the paragraph flow.
    # Appended to the LAST section as atomic blocks — a simplification (true
    # positional interleaving needs the underlying XML order), stated here
    # rather than hidden.
    if doc.tables and sections:
        for table in doc.tables:
            rows = ["| " + " | ".join(c.text.strip() for c in row.cells) + " |" for row in table.rows]
            sections[-1].atomic_blocks.append("\n".join(rows))
    return sections


# The dispatch table IS the format seam: one entry per supported suffix, and
# the Azure Document Intelligence connector is one more entry on the private
# side — nothing above this line changes when it arrives.
_PARSERS = {
    ".md": lambda data, doc_title: load_markdown(data.decode("utf-8", errors="replace"), doc_title=doc_title),
    ".pdf": load_pdf,
    ".docx": load_docx,
}

SUPPORTED_SUFFIXES = tuple(sorted(_PARSERS))


def parse_upload(filename: str, data: bytes, *, doc_title: str) -> list[Section]:
    """Uploaded bytes -> Sections, by file suffix. Raises ValueError for
    unsupported types — the route turns that into a 415, never a 500: a bad
    upload is the CLIENT's error."""
    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise ValueError(f"unsupported file type {suffix or '(none)'}; accepted: {', '.join(SUPPORTED_SUFFIXES)}")
    return parser(data, doc_title=doc_title)
